import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import io
import os
import numpy as np

def clean_covid_data():
    print("Processing Covid Data...")
    # Load dataset
    df = pd.read_csv('data visualisation/covid_analysis/COVID19_dataset.csv')
    
    # Drop rows with missing age or id
    df.dropna(subset=['age', 'id'], inplace=True)
    
    # Impute symptom_onset with mode
    mode_val = df['symptom_onset'].mode()[0]
    df['symptom_onset'] = df['symptom_onset'].fillna(mode_val)
    
    # Sort by date_onset_symptoms (asc) and age (desc)
    df.sort_values(by=['date_onset_symptoms', 'age'], ascending=[True, False], inplace=True)
    
    code = """
# Task 3: Covid Data Cleaning
df = pd.read_csv('data visualisation/covid_analysis/COVID19_dataset.csv')
df.dropna(subset=['age', 'id'], inplace=True)
mode_val = df['symptom_onset'].mode()[0]
df['symptom_onset'] = df['symptom_onset'].fillna(mode_val)
df.sort_values(by=['date_onset_symptoms', 'age'], ascending=[True, False], inplace=True)
    """
    return code, df.head().to_string()

def generate_sales_dashboard():
    print("Generating Sales Dashboard...")
    # Using mock data as sales_data.csv does not contain the required columns
    dates = pd.date_range(start='2023-01-01', end='2023-12-31', freq='D')
    countries = ['USA', 'UK', 'Canada', 'Germany', 'France', 'Australia']
    
    np.random.seed(42)
    data = {
        'Order Date': np.random.choice(dates, size=1000),
        'Country': np.random.choice(countries, size=1000),
        'Quantity': np.random.randint(1, 20, size=1000),
        'Unit Price': np.random.uniform(10, 1000, size=1000)
    }
    df = pd.DataFrame(data)
    df['Revenue'] = df['Quantity'] * df['Unit Price']
    
    # 1. Monthly Revenue (Line Plot)
    plt.figure(figsize=(6, 4))
    monthly_revenue = df.groupby(df['Order Date'].dt.to_period('M'))['Revenue'].sum()
    monthly_revenue.index = monthly_revenue.index.astype(str)
    plt.plot(monthly_revenue.index, monthly_revenue.values, marker='o')
    plt.title('Monthly Revenue Trend')
    plt.xlabel('Month')
    plt.ylabel('Revenue')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig('temp_sales_line.png')
    plt.close()
    
    # 2. Revenue by Country (Bar Plot)
    plt.figure(figsize=(6, 4))
    country_revenue = df.groupby('Country')['Revenue'].sum().sort_values(ascending=False)
    sns.barplot(x=country_revenue.index, y=country_revenue.values)
    plt.title('Revenue by Country')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig('temp_sales_bar.png')
    plt.close()
    
    # 3. Quantity Distribution (Histogram)
    plt.figure(figsize=(6, 4))
    sns.histplot(df['Quantity'], bins=20, kde=True)
    plt.title('Quantity Distribution')
    plt.tight_layout()
    plt.savefig('temp_sales_hist.png')
    plt.close()
    
    # 4. Price vs Revenue (Scatter Plot)
    plt.figure(figsize=(6, 4))
    sns.scatterplot(x='Unit Price', y='Revenue', data=df)
    plt.title('Price vs Revenue')
    plt.tight_layout()
    plt.savefig('temp_sales_scatter.png')
    plt.close()
    
    code = """
# Task 4: Sales Dashboard
# Note: Using mock data as sales_data.csv is incompatible
dates = pd.date_range(start='2023-01-01', end='2023-12-31', freq='D')
data = {
    'Order Date': np.random.choice(dates, size=1000),
    'Country': np.random.choice(['USA', 'UK', ...], size=1000),
    'Quantity': np.random.randint(1, 20, size=1000),
    'Unit Price': np.random.uniform(10, 1000, size=1000)
}
df = pd.DataFrame(data)
df['Revenue'] = df['Quantity'] * df['Unit Price']

# 1. Monthly Revenue
monthly_revenue = df.groupby(df['Order Date'].dt.to_period('M'))['Revenue'].sum()
plt.plot(monthly_revenue.index.astype(str), monthly_revenue.values)

# 2. Revenue by Country
sns.barplot(x='Country', y='Revenue', data=df, estimator=sum)

# 3. Quantity Distribution
sns.histplot(df['Quantity'], kde=True)

# 4. Price vs Revenue
sns.scatterplot(x='Unit Price', y='Revenue', data=df)
    """
    return code

def generate_iris_insights():
    print("Generating Iris Insights...")
    df = pd.read_csv('data visualisation/iris_flower/IRIS.csv')
    
    # 1. Pairplot
    plt.figure(figsize=(6, 6))
    sns.pairplot(df, hue='species')
    plt.savefig('temp_iris_pairplot.png')
    plt.close()
    
    # 2. Heatmap
    plt.figure(figsize=(6, 4))
    # Select only numeric columns for correlation matrix
    numeric_df = df.select_dtypes(include=[np.number])
    sns.heatmap(numeric_df.corr(), annot=True, cmap='coolwarm')
    plt.title('Correlation Heatmap')
    plt.tight_layout()
    plt.savefig('temp_iris_heatmap.png')
    plt.close()
    
    # 3. Boxplot
    plt.figure(figsize=(6, 4))
    sns.boxplot(x='species', y='petal_length', data=df)
    plt.title('Petal Length by Species')
    plt.tight_layout()
    plt.savefig('temp_iris_boxplot.png')
    plt.close()
    
    code = """
# Task 5: Iris Insights
# 1. Pairplot
sns.pairplot(df, hue='species')

# 2. Heatmap
sns.heatmap(df.select_dtypes(include=[np.number]).corr(), annot=True)

# 3. Boxplot
sns.boxplot(x='species', y='petal_length', data=df)
    """
    
    insights = """
    Insights:
    1. Petal length and petal width are highly correlated.
    2. Setosa is easily separable from the other two species based on petal dimensions.
    3. Virginica tends to have the largest petal length and width.
    """
    return code, insights

def create_document(covid_code, covid_output, sales_code, iris_code, iris_insights):
    print("Creating Word Document...")
    doc = Document()
    doc.add_heading('Data Analysis Report', 0)
    
    # Section 1: Healthcare Data Cleaning
    doc.add_heading('1. Healthcare Data Cleaning', level=1)
    doc.add_heading('Code:', level=2)
    doc.add_paragraph(covid_code)
    doc.add_heading('Output (Head):', level=2)
    doc.add_paragraph(covid_output)
    
    # Section 2: Sales Performance Dashboard
    doc.add_heading('2. Sales Performance Dashboard', level=1)
    doc.add_heading('Code:', level=2)
    doc.add_paragraph(sales_code)
    doc.add_heading('Visualizations:', level=2)
    
    doc.add_paragraph('Monthly Revenue Trend:')
    doc.add_picture('temp_sales_line.png', width=Inches(5))
    
    doc.add_paragraph('Revenue by Country:')
    doc.add_picture('temp_sales_bar.png', width=Inches(5))
    
    doc.add_paragraph('Quantity Distribution:')
    doc.add_picture('temp_sales_hist.png', width=Inches(5))
    
    doc.add_paragraph('Price vs Revenue:')
    doc.add_picture('temp_sales_scatter.png', width=Inches(5))
    
    # Section 3: Flower Classification Insights
    doc.add_heading('3. Flower Classification Insights', level=1)
    doc.add_heading('Code:', level=2)
    doc.add_paragraph(iris_code)
    doc.add_heading('Insights:', level=2)
    doc.add_paragraph(iris_insights)
    doc.add_heading('Visualizations:', level=2)
    
    doc.add_paragraph('Pairplot:')
    if os.path.exists('temp_iris_pairplot.png'):
        doc.add_picture('temp_iris_pairplot.png', width=Inches(5))
    
    doc.add_paragraph('Correlation Heatmap:')
    doc.add_picture('temp_iris_heatmap.png', width=Inches(5))
    
    doc.add_paragraph('Petal Length Boxplot:')
    doc.add_picture('temp_iris_boxplot.png', width=Inches(5))
    
    doc.save('Data_Analysis_Report.docx')
    print("Report saved as Data_Analysis_Report.docx")

def cleanup():
    files = [
        'temp_sales_line.png', 'temp_sales_bar.png', 'temp_sales_hist.png', 'temp_sales_scatter.png',
        'temp_iris_pairplot.png', 'temp_iris_heatmap.png', 'temp_iris_boxplot.png'
    ]
    for f in files:
        if os.path.exists(f):
            os.remove(f)
    print("Temporary files cleaned up.")

if __name__ == "__main__":
    try:
        covid_code, covid_output = clean_covid_data()
        sales_code = generate_sales_dashboard()
        iris_code, iris_insights = generate_iris_insights()
        
        create_document(covid_code, covid_output, sales_code, iris_code, iris_insights)
        cleanup()
    except Exception as e:
        print(f"An error occurred: {e}")
